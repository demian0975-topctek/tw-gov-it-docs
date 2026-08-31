#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
把 assets/工作說明書骨架.md 套上 assets/工作說明書範本.docx，產生 Word 草稿。

    python scripts/make_docx.py [骨架.md路徑] [輸出資料夾]

輸出檔名固定為 `工作說明書_V0.0.docx`；要保留舊版就在覆寫前自行備份。

## 版面規格只存在範本裡

字體、字級、縮排、行距、頁面邊界、頁首頁尾全部來自範本的 styles.xml 與 sectPr——
本腳本只負責判斷每一行是哪一層、掛上對應樣式，不寫任何版面數值。要調整版面請改範本
（`python tools/make_template.py <來源.docx>` 重新產生），改這裡改不到。

## 標號由 Word 自己算

壹／一／(一)／1／(1) 五層是範本裡一組多層清單，掛上樣式後 Word 自動編號，上層跳號時
下層歸 1。骨架 Markdown 裡寫的標號只是給人讀的，本腳本會**剝掉**再送進 Word，避免變成
「壹、壹、專案概述」。在 Word 裡插入或搬動章節後全篇自動重編，不需要回頭跑這支腳本。

表號與圖號則相反：在表格前一行單獨寫 `[表] 標題文字`（圖片用 `[圖]`），本腳本依當時
所屬章次算出「表 2-1」並寫成**文字**。文字表號不會跟著章節搬動而變，調整章節順序後
需重跑本腳本或於 Word 內手動修正。
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SKILL_DIR = Path(__file__).resolve().parent.parent
TEMPLATE = SKILL_DIR / "assets" / "工作說明書範本.docx"
SKELETON = SKILL_DIR / "assets" / "工作說明書骨架.md"
OUTPUT_NAME = "工作說明書_V0.0.docx"

# styleId 而非樣式名稱：範本沿用來源檔的 styleId（標題1 是 "10"、heading 3 是 "3"…），
# python-docx 會把 "heading 3" 這類內建名稱改寫成 "Heading 3" 而找不到樣式，掛 styleId 不會。
S_CHAPTER, S_SECTION, S_PAREN, S_ITEM, S_SUBITEM = "10", "2", "3", "4", "5"
S_BODY, S_COVER, S_COVER_NOTE, S_TABLE = "110", "cover", "coverNote", "tableText"

HEADER_PLACEHOLDER = "○○○　○○○"

HEADING_RE = re.compile(r"^(#{2,4})\s+(.*)")
CHAPTER_LABEL_RE = re.compile(r"^[壹貳參肆伍陸柒捌玖拾]+、\s*")
SECTION_LABEL_RE = re.compile(r"^[一二三四五六七八九十]+、\s*")
PAREN_RE = re.compile(r"^[（(][一二三四五六七八九十]+[)）]\s*(.*)")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)")
ORDERED_PAREN_RE = re.compile(r"^\(\d+\)\s*(.*)")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
BLOCKQUOTE_RE = re.compile(r"^>\s?")
CAPTION_RE = re.compile(r"^\[(表|圖)\]\s*(.*)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def styled(doc, style_id, text=""):
    """新增段落並掛上 styleId；行內 **粗體** 拆成獨立 run。"""
    p = doc.add_paragraph()
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    p._p.get_or_add_pPr().insert(0, pStyle)
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        p.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def fill_header(doc, agency, case_name):
    for paragraph in doc.sections[0].header.paragraphs:
        for run in paragraph.runs:
            if HEADER_PLACEHOLDER in run.text:
                run.text = f"{agency}　{case_name}"


def build_cover(doc, meta):
    for _ in range(4):
        styled(doc, S_COVER)
    styled(doc, S_COVER, meta["機關全銜"])
    styled(doc, S_COVER, meta["案名"])
    styled(doc, S_COVER, "工作說明書")
    for _ in range(6):
        styled(doc, S_COVER)
    styled(doc, S_COVER_NOTE, f"版本：{meta['版本']}")
    styled(doc, S_COVER_NOTE, f"日期：{meta['日期']}")
    styled(doc, S_COVER_NOTE, "機關內部工作稿，非正式版")
    doc.add_page_break()


def build_version_table(doc):
    # 版本紀錄是前置頁，掛封面樣式而非標題樣式——標題樣式帶自動編號，會被算成「一、」。
    styled(doc, S_COVER, "文件版本紀錄")
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for values, bold in ((["版本", "日期", "修訂內容", "修訂人"], True),
                         (["V0.0", "○○○", "顧問初稿", "○○○"], False)):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.paragraphs[0].style = doc.styles["表格內文"]
            cell.paragraphs[0].add_run(value).bold = bold
    doc.add_page_break()


class Renderer:
    """Markdown 每一行 → 骨架層級 → 範本樣式。層級判斷只看行首標號，不看縮排。"""

    def __init__(self, doc):
        self.doc = doc
        self.chapter_no = 0
        self.caption_seq = {}

    def render(self, md_text):
        table_buf = []
        for raw in md_text.split("\n"):
            line = raw.strip()

            if TABLE_ROW_RE.match(line):
                table_buf.append(line)
                continue
            if table_buf:
                self._flush_table(table_buf)
                table_buf = []
            if not line or BLOCKQUOTE_RE.match(line):
                # > 區塊是骨架給撰寫者看的挖空與寫法指引，不進入交付的 Word 檔。
                continue
            self._render_line(line)

        if table_buf:
            self._flush_table(table_buf)

    def _render_line(self, line):
        m = HEADING_RE.match(line)
        if m:
            level, text = len(m.group(1)), m.group(2).strip()
            if level == 2:
                self.chapter_no += 1
                styled(self.doc, S_CHAPTER, CHAPTER_LABEL_RE.sub("", text))
            elif level == 3:
                styled(self.doc, S_SECTION, SECTION_LABEL_RE.sub("", text))
            else:
                styled(self.doc, S_PAREN, PAREN_RE.sub(r"\1", text))
            return

        m = CAPTION_RE.match(line)
        if m:
            self._add_caption(m.group(1), m.group(2))
            return

        for pattern, style_id in ((PAREN_RE, S_PAREN),
                                  (ORDERED_RE, S_ITEM),
                                  (ORDERED_PAREN_RE, S_SUBITEM)):
            m = pattern.match(line)
            if m:
                styled(self.doc, style_id, m.group(1))
                return

        styled(self.doc, S_BODY, line)

    def _add_caption(self, kind, text):
        key = (kind, self.chapter_no)
        self.caption_seq[key] = self.caption_seq.get(key, 0) + 1
        styled(self.doc, S_TABLE, f"{kind} {self.chapter_no}-{self.caption_seq[key]}　{text}")

    def _flush_table(self, rows_raw):
        rows = [[c.strip() for c in r.strip("|").split("|")] for r in rows_raw]
        if len(rows) >= 2 and set(rows[1][0]) <= set("-: "):
            del rows[1]
        table = self.doc.add_table(rows=0, cols=len(rows[0]))
        table.style = "Table Grid"
        for r_i, row in enumerate(rows):
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.paragraphs[0].style = self.doc.styles["表格內文"]
                cell.paragraphs[0].add_run(value).bold = (r_i == 0)


def read_meta(meta_block):
    meta = {"機關全銜": "○○○", "案名": "○○○", "承辦單位": "○○○",
            "版本": "顧問草稿 V0.0", "日期": "○○○"}
    for key, val in re.findall(r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|", meta_block):
        if key in meta and "---" not in val:
            meta[key] = val
    return meta


def main():
    md_path = Path(sys.argv[1]) if len(sys.argv) > 1 else SKELETON
    out_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else Path.cwd()

    text = md_path.read_text(encoding="utf-8")
    body_start = text.index("## 壹、")
    meta = read_meta(text[:body_start])

    doc = Document(TEMPLATE)
    fill_header(doc, meta["機關全銜"], meta["案名"])
    build_cover(doc, meta)
    build_version_table(doc)
    Renderer(doc).render(text[body_start:])

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / OUTPUT_NAME
    doc.save(out_path)
    print(f"已產生：{out_path}")


if __name__ == "__main__":
    main()
