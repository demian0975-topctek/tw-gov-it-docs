#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
從 assets/ 的骨架 Markdown 產生符合公部門版面慣例的 Word 空白文件。

用法：
    pip install python-docx

    # 直接以參數帶入
    python scripts/make_docx_skeleton.py 03 --機關 "○○部" --案名 "○○系統建置案" \
        --廠商 "○○資訊股份有限公司" -o 需求規格書.docx

    # 以 JSON 帶入專案基本資料（同時給參數時，參數優先）
    python scripts/make_docx_skeleton.py 03 --json 專案.json

    # 一次產生全部骨架到指定目錄
    python scripts/make_docx_skeleton.py --all --json 專案.json --outdir out/

    python scripts/make_docx_skeleton.py --list       # 列出可用骨架
    python scripts/make_docx_skeleton.py --範例json   # 印出 JSON 範本

JSON 格式（欄位皆可省略）：
    {
      "機關": "○○部",
      "案名": "○○系統建置案",
      "廠商": "○○資訊股份有限公司",
      "版本": "v1.0",
      "日期": "○○○ 年 ○○ 月 ○○ 日",
      "文件編號": "ABC-SRS-001"
    }

版面設定（各機關要求不同，可用參數覆寫）：
    A4、上下 2.54cm 左右 3.17cm、內文標楷體 12pt、固定行高 20pt
    頁首置中「機關全銜　案名」、頁尾置中頁碼
    含封面、文件版本紀錄表、可自動更新的目錄與表目錄，以及圖目錄的位置與填法指引

注意：目錄、表目錄與表格標號以 Word 功能變數產生，開啟檔案後需按 Ctrl+A 再按 F9
      （macOS 為 Cmd+A 後於「參考資料」更新目錄）才會顯示內容。
      **圖目錄例外**：骨架尚無任何「圖」標號，預先插入功能變數在 Word 中會顯示
      「找不到圖表目錄項目」，故該處放的是【待填】指引，按 F9 不會有變化；
      待文件插入圖片並建立「圖」標號後，再依指引自行插入圖表目錄。
      表格標號採流水號（表 1、表 2…）。若機關要求「表 3-1」的章-序格式，
      因本文採中文數字章名，無法以功能變數自動產生，需於 Word 內自行調整。
"""
import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    sys.exit("需要 python-docx，請先執行：pip install python-docx")

ASSETS = Path(__file__).resolve().parent.parent / "assets"

BODY_FONT = "標楷體"
HEAD_FONT = "標楷體"
ASCII_FONT = "Times New Roman"

DEFAULTS = {
    "機關": "○○○○○○（機關全銜）",
    "案名": "○○○○○○○○（案名）",
    "廠商": "○○○○股份有限公司",
    "版本": "v1.0",
    "日期": "○○○ 年 ○○ 月 ○○ 日",
    "文件編號": "【待填】",
}


def set_font(run, name=BODY_FONT, size=12, bold=False):
    run.font.name = ASCII_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)


def add_para(doc, text="", size=12, bold=False, align=None, indent=0,
             font=BODY_FONT, space_after=0, outline_level=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    if align is not None:
        p.alignment = align
    if outline_level is not None:
        # 目錄功能變數靠 outlineLvl 認段落，沒有這一段目錄會是空的
        ppr = p._element.get_or_add_pPr()
        el = OxmlElement("w:outlineLvl")
        el.set(qn("w:val"), str(outline_level))
        ppr.append(el)
    if text:
        set_font(p.add_run(text), font, size, bold)
    return p


def add_field(paragraph, instruction, placeholder="", size=12, font=BODY_FONT):
    """插入 Word 功能變數（TOC、PAGE、SEQ 等）。"""
    run = paragraph.add_run()
    set_font(run, font, size)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, sep):
        run._element.append(el)
    if placeholder:
        ph = OxmlElement("w:t")
        ph.text = placeholder
        run._element.append(ph)
    run._element.append(end)
    return run


def build_header_footer(section, agency, project):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(hp.add_run(f"{agency}　{project}"), BODY_FONT, 10)
    bottom_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, " PAGE ", "1", size=10)


def bottom_border(paragraph):
    ppr = paragraph._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    borders.append(bottom)
    ppr.append(borders)


def build_cover(doc, meta, doc_title):
    for _ in range(4):
        add_para(doc)
    add_para(doc, meta["機關"], size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, meta["案名"], size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        add_para(doc)
    add_para(doc, doc_title, size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6):
        add_para(doc)
    for line in (f"版　　本：{meta['版本']}",
                 f"文件編號：{meta['文件編號']}",
                 f"承辦廠商：{meta['廠商']}",
                 f"中華民國 {meta['日期']}"):
        add_para(doc, line, size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()


def build_version_table(doc):
    add_para(doc, "文件版本紀錄", size=16, bold=True, font=HEAD_FONT, space_after=6)
    headers = ["版本", "日期", "修訂內容摘要", "修訂人", "審核人"]
    t = doc.add_table(rows=4, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.text = ""
        set_font(cell.paragraphs[0].add_run(h), BODY_FONT, 12, True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(t.cell(1, 0).paragraphs[0].add_run("v0.1"), BODY_FONT, 11)
    doc.add_page_break()


def build_toc(doc):
    for title, instruction in (
        ("目　　錄", r' TOC \o "1-3" \h \z \u '),
        ("表　目　錄", r' TOC \h \z \c "表" '),
    ):
        add_para(doc, title, size=18, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, font=HEAD_FONT, space_after=6)
        p = add_para(doc)
        add_field(p, instruction, "【請於 Word 中按 F9 更新目錄】", size=11)
        doc.add_page_break()

    # 圖目錄不預先插入功能變數：骨架尚無任何「圖」標號，空的圖表目錄
    # 在 Word 中會直接顯示「找不到圖表目錄項目」的錯誤字串。
    add_para(doc, "圖　目　錄", size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, font=HEAD_FONT, space_after=6)
    add_para(doc, "【待填：插入圖片並以「參考資料 → 插入標號」建立「圖」標號後，"
                  "於此處執行「參考資料 → 插入圖表目錄」，標籤選「圖」】",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_caption(doc, kind, text="【待填：標題】"):
    """插入「表 { SEQ 表 }　標題」標號段落，供圖表目錄擷取。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = doc.styles["Caption"]
    except KeyError:
        pass
    set_font(p.add_run(f"{kind} "), BODY_FONT, 11, True)
    add_field(p, f' SEQ {kind} \\* ARABIC ', "1", size=11)
    set_font(p.add_run(f"　{text}"), BODY_FONT, 11)
    return p


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
SEP_ROW_RE = re.compile(r"^\|[\s:|-]+\|\s*$")
HR_RE = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")  # 分隔線不進 Word
# 骨架中的 ### 一、xxx 這類段落，Word 目錄取到第 3 層即可
OUTLINE_BY_LEVEL = {2: 0, 3: 1, 4: 2, 5: 3, 6: 4}
HEAD_SIZE_BY_LEVEL = {2: 16, 3: 14, 4: 13, 5: 12, 6: 12}


def render_markdown(doc, md_text, captions=True):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            add_para(doc, line, size=11, indent=1, font="Consolas")
            i += 1
            continue

        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                i += 1
                continue  # 標題已在封面呈現
            add_para(doc, text, size=HEAD_SIZE_BY_LEVEL.get(level, 12), bold=True,
                     font=HEAD_FONT, space_after=6,
                     outline_level=OUTLINE_BY_LEVEL.get(level, 4))
            i += 1
            continue

        if TABLE_ROW_RE.match(line):
            block = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i].rstrip()):
                if not SEP_ROW_RE.match(lines[i].rstrip()):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    block.append(cells)
                i += 1
            if block:
                if captions:
                    add_caption(doc, "表")
                cols = max(len(r) for r in block)
                rows = max(len(block), 3)
                t = doc.add_table(rows=rows, cols=cols)
                t.style = "Table Grid"
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r, row in enumerate(block):
                    for c, val in enumerate(row):
                        cell = t.cell(r, c)
                        cell.text = ""
                        set_font(cell.paragraphs[0].add_run(val), BODY_FONT, 11, r == 0)
                        if r == 0:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_para(doc)
            continue

        if line.startswith(">"):
            add_para(doc, line.lstrip("> ").strip(), size=11, indent=1)
            i += 1
            continue

        if not line or HR_RE.match(line):
            i += 1
            continue

        add_para(doc, line)
        i += 1


def build_document(src: Path, meta: dict, captions=True) -> tuple:
    title = re.sub(r"^\d+-", "", src.stem).replace("骨架", "").strip()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.17)
    build_header_footer(sec, meta["機關"], meta["案名"])

    build_cover(doc, meta, title)
    build_version_table(doc)
    build_toc(doc)
    render_markdown(doc, src.read_text(encoding="utf-8"), captions=captions)
    return doc, title


SAMPLE_JSON = json.dumps(DEFAULTS, ensure_ascii=False, indent=2)


def main():
    ap = argparse.ArgumentParser(add_help=True)
    ap.add_argument("skeleton", nargs="?", help="骨架編號或檔名關鍵字，例如 03 或 需求規格書")
    ap.add_argument("--list", action="store_true", help="列出可用骨架")
    ap.add_argument("--all", action="store_true", help="產生全部骨架")
    ap.add_argument("--outdir", default=".", help="--all 時的輸出目錄")
    ap.add_argument("--json", dest="json_path", help="以 JSON 檔帶入專案基本資料")
    ap.add_argument("--範例json", dest="sample", action="store_true", help="印出 JSON 範本")
    ap.add_argument("--無標號", dest="no_caption", action="store_true",
                    help="不自動插入表格標號")
    ap.add_argument("--機關", dest="agency")
    ap.add_argument("--案名", dest="project")
    ap.add_argument("--廠商", dest="vendor")
    ap.add_argument("--版本", dest="version")
    ap.add_argument("--日期", dest="date")
    ap.add_argument("--文件編號", dest="docno")
    ap.add_argument("-o", "--output", default=None)
    args = ap.parse_args()

    if args.sample:
        print(SAMPLE_JSON)
        return

    files = sorted(ASSETS.glob("*.md"))
    if args.list or (not args.skeleton and not args.all):
        print("可用骨架：")
        for f in files:
            print(f"  {f.stem}")
        return

    meta = dict(DEFAULTS)
    if args.json_path:
        raw = json.loads(Path(args.json_path).read_text(encoding="utf-8"))
        unknown = set(raw) - set(DEFAULTS)
        if unknown:
            print(f"警告：JSON 中有未使用的欄位：{'、'.join(sorted(unknown))}", file=sys.stderr)
        meta.update({k: str(v) for k, v in raw.items() if k in DEFAULTS})
    for key, val in (("機關", args.agency), ("案名", args.project), ("廠商", args.vendor),
                     ("版本", args.version), ("日期", args.date), ("文件編號", args.docno)):
        if val:
            meta[key] = val

    captions = not args.no_caption

    if args.all:
        outdir = Path(args.outdir)
        outdir.mkdir(parents=True, exist_ok=True)
        for src in files:
            doc, title = build_document(src, meta, captions)
            out = outdir / f"{src.stem}.docx"
            doc.save(out)
            print(f"已產生：{out}")
        print("\n提示：開啟檔案後按 Ctrl+A、F9 更新目錄與標號。")
        return

    matches = [f for f in files if args.skeleton in f.stem]
    if not matches:
        sys.exit(f"找不到骨架：{args.skeleton}（用 --list 查看）")
    src = matches[0]

    doc, title = build_document(src, meta, captions)
    out = args.output or f"{title}.docx"
    doc.save(out)
    print(f"已產生：{out}（來源骨架：{src.name}）")
    print("提示：開啟檔案後按 Ctrl+A、F9 更新目錄與標號。")


if __name__ == "__main__":
    main()
