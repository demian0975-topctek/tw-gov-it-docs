#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
比對 .docx 的版面與「政府文書格式參考規範」第四點，用於產出後自我驗證。

用法：
    pip install python-docx

    python scripts/check_layout.py 需求規格書.docx
    python scripts/check_layout.py out/*.docx          # 可一次驗多份
    python scripts/check_layout.py --只看錯誤 *.docx   # 只列不符項
    python scripts/check_layout.py --rules             # 列出目前載入的判定值

    # 版面來自機關範本時：改比對範本，不比對規範值
    python scripts/check_layout.py --範本 機關範本.docx 產出.docx

離開碼：0 全部符合；1 有不符項；2 用法或環境錯誤。

兩個模式的差別是**拿什麼當基準**：預設模式比對規範值，適用於版面由本 skill 自己產生的
文件；`--範本` 比對指定的 .docx，適用於版面住在機關範本裡的文件——實際案件的範本常見
四邊 2.0 公分，與公文格式本來就不同，用預設模式驗只會得到滿螢幕假警報。`--範本` 驗的是產出有沒有偏離範本，也就是腳本有沒有偷偷蓋掉範本的版面。

判定值不在本檔，在 references/06-Word範本.md 的【版面檢查值】表格——要改
判定標準改那張表。這與 check_wording.py、check_sources.py 的做法一致。

**驗的是「有效值」，不是「有沒有明確設定」。** Word 的格式沿
「run → 段落 → 段落樣式 → Normal」繼承，直接數 run 上的屬性，會把從樣式正確
繼承來的值誤判成未設定；空白表格列與空段落更是完全沒有 run 可數。這一點寫在
這裡是因為第一版腳本就是這樣寫的，修好 Normal 樣式後數字卻毫無變化——量錯了東西。

**該規範第三點的適用範圍是公文與證書獎狀，技術文件不在其內。** 本腳本是給
「機關要求比照公文格式」時用的；機關另有範本時以機關範本為準，不要拿這裡的
NG 去否定機關的版面。適用範圍見 references/06-Word範本.md 第二節。
"""
import argparse
import re
import sys
from collections import Counter
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("需要 python-docx，請先執行：pip install python-docx")

SPEC_DOC = Path(__file__).resolve().parent.parent / "references" / "06-Word範本.md"
SPEC_HEADING = "【版面檢查值】"
EMU_CM = 360000
ROUNDING_TOL = 0.01  # Cm() 換算成 EMU 的捨入誤差，不是規範的容許值

RANGE_RE = re.compile(r"([\d.]+)\s*～\s*([\d.]+)\s*(cm|pt)")


def load_spec(path=SPEC_DOC):
    """解析【版面檢查值】表格。規則住在表格裡，這裡只負責讀。"""
    if not path.exists():
        sys.exit(f"找不到判定值來源：{path}")
    lines = path.read_text(encoding="utf-8").splitlines()
    try:
        start = next(i for i, ln in enumerate(lines) if SPEC_HEADING in ln)
    except StopIteration:
        sys.exit(f"{path} 裡找不到「{SPEC_HEADING}」表格")

    ranges, fonts = {}, {}
    for ln in lines[start:]:
        if not ln.startswith("|"):
            if ranges or fonts:      # 表格已結束
                break
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cells) < 3 or cells[1] in ("型態", "---"):
            continue
        item, kind, allowed = cells[0], cells[1], cells[2]
        if kind == "範圍":
            m = RANGE_RE.search(allowed)
            if m:
                ranges[item] = (float(m.group(1)), float(m.group(2)), m.group(3))
        elif kind == "字型":
            fonts[item] = [v.strip() for v in allowed.split("、") if v.strip()]
    if not ranges or not fonts:
        sys.exit(f"{path} 的【版面檢查值】表格解析不出內容")
    return ranges, fonts


def style_chain(style):
    """由段落樣式往上走到最上層，回傳整條繼承鏈。"""
    seen = []
    while style is not None and style not in seen:
        seen.append(style)
        style = style.base_style
    return seen


def doc_default_font(doc, attr):
    """styles.xml 的 w:docDefaults——樣式鏈全部落空後 Word 真正會用的字型。

    少了這一層會把「繼承自文件預設」誤報成「未繼承到任何字型」，而文件預設
    往往正是新細明體或 Calibri，也就是最該被抓出來的那個值。
    """
    dd = doc.styles.element.find(qn("w:docDefaults"))
    rpd = dd.find(qn("w:rPrDefault")) if dd is not None else None
    rpr = rpd.find(qn("w:rPr")) if rpd is not None else None
    f = rpr.find(qn("w:rFonts")) if rpr is not None else None
    if f is None:
        return None
    val = f.get(qn(f"w:{attr}"))
    if val:
        return val
    theme = f.get(qn(f"w:{attr}Theme"))
    return f"（佈景主題 {theme}）" if theme else None


def doc_default_spacing(doc):
    """w:docDefaults 的段落行距。lineRule=auto 時 w:line 是 240 分之一行，
    非點值——換算成倍數回傳，才不會把 1.15 倍當成 276 點。"""
    dd = doc.styles.element.find(qn("w:docDefaults"))
    ppd = dd.find(qn("w:pPrDefault")) if dd is not None else None
    ppr = ppd.find(qn("w:pPr")) if ppd is not None else None
    sp = ppr.find(qn("w:spacing")) if ppr is not None else None
    if sp is None:
        return None, None
    line = sp.get(qn("w:line"))
    if line is None:
        return None, None
    rule = sp.get(qn("w:lineRule"), "auto")
    if rule == "auto":
        return round(int(line) / 240, 2), WD_LINE_SPACING.MULTIPLE
    return round(int(line) / 20, 1), (WD_LINE_SPACING.EXACTLY if rule == "exact"
                                      else WD_LINE_SPACING.AT_LEAST)


def eff_font(run, para, doc, attr):
    """run 實際生效的字型：run → 段落樣式鏈 → Normal → docDefaults。"""
    rpr = run._element.rPr
    f = rpr.find(qn("w:rFonts")) if rpr is not None else None
    if f is not None and f.get(qn(f"w:{attr}")):
        return f.get(qn(f"w:{attr}"))
    for st in style_chain(para.style) + [doc.styles["Normal"]]:
        rpr = st.element.rPr
        f = rpr.find(qn("w:rFonts")) if rpr is not None else None
        if f is not None and f.get(qn(f"w:{attr}")):
            return f.get(qn(f"w:{attr}"))
    return doc_default_font(doc, attr)


def eff_spacing(para, doc):
    """段落實際生效的行距與規則：段落 → 樣式鏈 → Normal → docDefaults。"""
    pf = para.paragraph_format
    if pf.line_spacing is not None:
        return pf.line_spacing, pf.line_spacing_rule
    for st in style_chain(para.style) + [doc.styles["Normal"]]:
        spf = st.paragraph_format
        if spf.line_spacing is not None:
            return spf.line_spacing, spf.line_spacing_rule
    return doc_default_spacing(doc)


def all_paragraphs(doc):
    """本文加表格儲存格。表格是使用者實際填字的地方，不能略過。"""
    out = [(p, "本文") for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                out += [(p, "表格") for p in cell.paragraphs]
    return out


def style_fingerprint(doc):
    """每個樣式的版面指紋，用來比對產出與範本有無走樣。"""
    out = {}
    for st in doc.styles:
        try:
            pf = st.paragraph_format
        except (AttributeError, ValueError):
            continue          # 字元樣式沒有段落格式
        rpr = st.element.rPr
        f = rpr.find(qn("w:rFonts")) if rpr is not None else None
        ls = pf.line_spacing
        out[st.name] = (
            f.get(qn("w:eastAsia")) if f is not None else None,
            f.get(qn("w:ascii")) if f is not None else None,
            round(ls.pt, 1) if hasattr(ls, "pt") else ls,
            str(pf.line_spacing_rule),
        )
    return out


def check_against_template(path, tpl_path, quiet=False):
    """比對產出與 .docx 範本，而非比對「政府文書格式參考規範」。

    範本由真實案件萃取，版面本來就可能與公文格式不同（本案範本即為四邊 2.0
    公分），拿規範去驗只會得到滿screen假警報。這裡驗的是 CLAUDE.md 真正在意
    的事：**腳本有沒有偷偷蓋掉範本的版面**——那會變成兩份互相矛盾的版面規格，
    而 Word 開檔時以範本為準，腳本裡那份永遠看不出有沒有生效。
    """
    doc, tpl = Document(path), Document(tpl_path)
    ng, lines = [], []

    s, ts = doc.sections[0], tpl.sections[0]
    for item, attr in (("頁面寬", "page_width"), ("頁面高", "page_height"),
                       ("上緣", "top_margin"), ("下緣", "bottom_margin"),
                       ("左緣", "left_margin"), ("右緣", "right_margin")):
        got = getattr(s, attr)
        want = getattr(ts, attr)
        got = None if got is None else round(got / EMU_CM, 3)
        want = None if want is None else round(want / EMU_CM, 3)
        ok = got is not None and want is not None and abs(got - want) <= ROUNDING_TOL
        if not ok:
            ng.append(f"{item} {got}≠範本 {want} cm")
        lines.append(f"  {'OK ' if ok else 'NG '} {item:<6} {got} cm"
                     f"（範本 {want}）")

    got_st, want_st = style_fingerprint(doc), style_fingerprint(tpl)
    drifted = [n for n in got_st if n in want_st and got_st[n] != want_st[n]]
    extra = [n for n in got_st if n not in want_st]
    if drifted:
        ng.append(f"{len(drifted)} 個樣式與範本不同")
    if extra:
        ng.append(f"{len(extra)} 個樣式範本沒有")
    lines.append(f"  {'OK ' if not drifted else 'NG '} 樣式定義 "
                 f"{len(got_st)} 個，與範本相同 {len(got_st) - len(drifted) - len(extra)} 個")
    for n in drifted[:5]:
        lines.append(f"       走樣：{n} {got_st[n]} ≠ 範本 {want_st[n]}")
    for n in extra[:5]:
        lines.append(f"       範本沒有的樣式：{n}")

    # 繼承不到行距或字型的段落。範本自己的樣式就沒定義時（本案範本的標題1、
    # 標題2 即是，真實案件本來就那樣寫），那是機關的設計而非產出走樣——只在
    # 樣式是範本沒有的、也就是真的走樣時才判 NG，其餘列為提示。
    orphan = Counter()
    for para, where in all_paragraphs(doc):
        if eff_spacing(para, doc)[0] is not None and not any(
                eff_font(r, para, doc, a) is None
                for r in para.runs for a in ("eastAsia", "ascii")):
            continue
        orphan[para.style.name] += 1
    drift_orphan = {n: c for n, c in orphan.items() if n in extra}
    inherited = {n: c for n, c in orphan.items() if n not in extra}
    if drift_orphan:
        ng.append(f"{sum(drift_orphan.values())} 段落在範本沒有的樣式上且繼承不到格式")
    lines.append(f"  {'OK ' if not drift_orphan else 'NG '} "
                 f"無「範本沒有的樣式」造成的孤兒段落")
    for n, c in inherited.items():
        lines.append(f"       提示：樣式 {n} 未定義行距（範本即如此，非產出走樣）"
                     f"，{c} 段沿用 Word 預設行距")

    if not quiet or ng:
        print(f"\n=== {Path(path).name}（對照範本 {Path(tpl_path).name}）===")
        for ln in lines:
            if not quiet or ln.lstrip().startswith(("NG", "走樣", "範本沒有的樣式")):
                print(ln)
        print(f"  結論：{'與範本一致' if not ng else '不一致：' + '；'.join(ng)}")
    return not ng


def check(path, ranges, fonts, quiet=False):
    doc = Document(path)
    sec = doc.sections[0]
    ng = []

    def rng(item, got):
        lo, hi, unit = ranges[item]
        ok = got is not None and lo - ROUNDING_TOL <= got <= hi + ROUNDING_TOL
        if not ok:
            ng.append(f"{item} {got} {unit}（規範 {lo}～{hi}）")
        return ok, f"{got} {unit}", f"{lo}～{hi}"

    lines = []
    for item, got in (("頁面寬", sec.page_width), ("頁面高", sec.page_height),
                      ("上緣", sec.top_margin), ("下緣", sec.bottom_margin),
                      ("左緣", sec.left_margin), ("右緣", sec.right_margin)):
        ok, shown, spec = rng(item, None if got is None else round(got / EMU_CM, 3))
        lines.append(f"  {'OK ' if ok else 'NG '} {item:<6} {shown:<14}（{spec}）")

    paras = all_paragraphs(doc)
    lo, hi, _ = ranges["行距"]
    spacing, bad_space = Counter(), []
    for para, where in paras:
        ls, rule = eff_spacing(para, doc)
        if ls is None:
            bad_space.append((where, para.style.name, "無行距"))
            spacing[("無行距", where)] += 1
            continue
        # 只有固定行高（EXACTLY／AT_LEAST）的值才是點數；倍數行高的 1.15
        # 是「1.15 倍」，標成 pt 會讓人以為是 1.15 點而完全看不出問題。
        fixed = rule in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST)
        v = round(ls.pt, 1) if hasattr(ls, "pt") else ls
        shown = f"{v} pt" if fixed else f"{v} 倍"
        # 規範講的是固定行高的點數，所以倍數行高一律不符，不論倍數多少。
        ok = rule == WD_LINE_SPACING.EXACTLY and lo <= v <= hi
        spacing[(shown, where, ok)] += 1
        if not ok:
            bad_space.append((where, para.style.name, shown))
    for (shown, where, ok), n in sorted(spacing.items(), key=lambda kv: str(kv[0])):
        lines.append(f"  {'OK ' if ok else 'NG '} 行距   {shown}（{where}）：{n} 段")
    if bad_space:
        ng.append(f"行距不符 {len(bad_space)} 段")

    allowed_ea = set(fonts["中文字型"]) | set(fonts["等寬字型例外"])
    allowed_asc = set(fonts["英數字型"]) | set(fonts["等寬字型例外"])
    ea, asc = Counter(), Counter()
    for para, _ in paras:
        for run in para.runs:
            ea[eff_font(run, para, doc, "eastAsia")] += 1
            asc[eff_font(run, para, doc, "ascii")] += 1
    for label, got, allowed in (("中文字型", ea, allowed_ea),
                                ("英數字型", asc, allowed_asc)):
        bad = set(got) - allowed
        if bad:
            ng.append(f"{label}出現 "
                      + "、".join(b if b else "未繼承到任何字型" for b in bad))
        lines.append(f"  {'OK ' if not bad else 'NG '} {label} "
                     + "、".join(f"{k or '未繼承到任何字型'}×{n}"
                                 for k, n in got.most_common()))

    if not quiet or ng:
        print(f"\n=== {Path(path).name} ===")
        for ln in lines:
            if not quiet or ln.lstrip().startswith("NG"):
                print(ln)
        print(f"  結論：{'符合規範第四點' if not ng else '不符：' + '；'.join(ng)}"
              f"（共 {len(paras)} 段）")
        for where, style, v in bad_space[:5]:
            print(f"       行距不符：{where} 樣式 {style} → {v}")
    return not ng


def main():
    ap = argparse.ArgumentParser(
        description="比對 .docx 版面與「政府文書格式參考規範」第四點")
    ap.add_argument("paths", nargs="*", help="要檢查的 .docx")
    ap.add_argument("--只看錯誤", dest="quiet", action="store_true",
                    help="只列不符項")
    ap.add_argument("--範本", dest="template", metavar="範本.docx",
                    help="改為比對這份 .docx 範本，而非比對規範值。"
                         "版面住在範本裡的文件用這個模式")
    ap.add_argument("--rules", action="store_true",
                    help="列出目前載入的判定值後結束")
    args = ap.parse_args()

    if args.template:
        if not args.paths:
            ap.error("請指定要檢查的 .docx")
        if not Path(args.template).exists():
            sys.exit(f"找不到範本：{args.template}")
        results = [check_against_template(p, args.template, args.quiet)
                   for p in args.paths]
        bad = results.count(False)
        print(f"\n合計：{len(results)} 份，與範本一致 {results.count(True)} 份、"
              f"不一致 {bad} 份。")
        if bad:
            print("版面規格只存在範本裡。要改版面改範本後重跑 "
                  "tools/make_template.py，不要在腳本裡寫 Pt/Twips 數值。")
        return 1 if bad else 0

    ranges, fonts = load_spec()
    if args.rules:
        print(f"判定值來源：{SPEC_DOC.name} 的{SPEC_HEADING}表格\n")
        for item, (lo, hi, unit) in ranges.items():
            print(f"  {item:<6} 範圍 {lo}～{hi} {unit}")
        for item, vals in fonts.items():
            print(f"  {item:<6} 字型 {'、'.join(vals)}")
        return 0

    if not args.paths:
        ap.error("請指定要檢查的 .docx（或用 --rules 列出判定值）")

    results = [check(p, ranges, fonts, args.quiet) for p in args.paths]
    bad = results.count(False)
    print(f"\n合計：{len(results)} 份，符合 {results.count(True)} 份、"
          f"不符 {bad} 份。")
    if bad:
        print("機關另有範本時以機關範本為準；本規範的範圍是公文與證書獎狀，"
              "適用判斷見 references/06-Word範本.md 第二節。")
    return 1 if bad else 0


if __name__ == "__main__":
    sys.exit(main())
